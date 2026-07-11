import os
import docx

source_dir = r"C:\Users\kasim\OneDrive\Masaüstü\dök"
target_dir = r"C:\Users\kasim\OneDrive\Masaüstü\Yeni klasör\public\templates"

if not os.path.exists(target_dir):
    os.makedirs(target_dir)

def process_docs():
    for root, _, files in os.walk(source_dir):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'): # ignore temp files like ~$file.docx
                file_path = os.path.join(root, file)
                target_path = os.path.join(target_dir, file)
                
                try:
                    doc = docx.Document(file_path)
                    text_to_insert = (
                        "[SİSTEM BİLGİSİ - ZORUNLU ALANLAR]\n"
                        "Firma: {companyName}\n"
                        "Tarih: {date}\n"
                        "Hazırlayan: {preparedBy}\n"
                        "Açıklama: {description}"
                    )
                    
                    if len(doc.paragraphs) > 0:
                        doc.paragraphs[0].insert_paragraph_before(text_to_insert)
                    else:
                        doc.add_paragraph(text_to_insert)
                        
                    doc.save(target_path)
                    print(f"Processed and saved: {file}")
                except Exception as e:
                    print(f"Failed to process {file}: {e}")

if __name__ == "__main__":
    process_docs()